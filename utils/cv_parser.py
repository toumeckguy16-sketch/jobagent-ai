"""Module CV Parser
Lit un CV uploadé (PDF ou DOCX) et en extrait un profil structuré
en utilisant le modèle Llama via l'API Groq (gratuit et rapide).
Formats supportés : PDF, DOCX, DOC
API Groq : https://console.groq.com  (gratuit, sans carte bancaire)
Modèle   : llama-3.3-70b-versatile"""
import os
import io
import re
import json
from typing import Optional
from pathlib import Path
# Extraction texte PDF
import pdfplumber
# Extraction texte DOCX
from docx import Document as DocxDocument
# Llama via Groq (LangChain)
from langchain_groq import ChatGroq
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from pydantic import BaseModel, Field
from typing import List
# ─────────────────────────────────────────────
#  SCHÉMA DU PROFIL EXTRAIT (Pydantic)
# ─────────────────────────────────────────────
class ProfessionalExperience(BaseModel):
    title: str = Field(description="Poste occupé")
    company: str = Field(description="Nom de l'entreprise")
    period: str = Field(description="Période ou durée (ex: 2020-2022, 6 mois)")
    description: str = Field(description="Courte description des tâches et réalisations")
    technologies: List[str] = Field(description="Technologies ou outils utilisés durant cette expérience")

class CandidateProfile(BaseModel):
    """Profil structuré extrait du CV par Llama"""
    full_name:        str       = Field(description="Nom complet du candidat")
    email:            str       = Field(description="Adresse email (vide si absente)")
    phone:            str       = Field(description="Numéro de téléphone (vide si absent)")
    location:         str       = Field(description="Ville / pays du candidat")
    job_title:        str       = Field(description="Titre du poste actuel ou recherché")
    summary:          str       = Field(description="Résumé du profil en 2-3 phrases")
    hard_skills:      List[str] = Field(description="Compétences techniques")
    soft_skills:      List[str] = Field(description="Compétences comportementales")
    tools:            List[str] = Field(description="Outils et logiciels maîtrisés")
    languages:        List[str] = Field(description="Langues parlées")
    education_level:  str       = Field(description="Niveau d'éducation le plus élevé")
    experience_years: int       = Field(description="Années d'expérience totales estimées")
    sectors:          List[str] = Field(description="Secteurs ciblés ou expérimentés")
    experiences:      List[ProfessionalExperience] = Field(description="Liste des expériences professionnelles")
    profile_text:     str       = Field(description="Paragraphe résumé utilisé pour la recherche d'offres")
# ─────────────────────────────────────────────
#  CV PARSER — Llama via Groq
# ─────────────────────────────────────────────
class CVParser:
    """
    Lit un CV (PDF ou DOCX) et extrait un profil structuré
    en utilisant Llama-3.3-70b via l'API Groq.
    Pourquoi Groq + Llama ?
    - API 100% gratuite (pas de carte bancaire requise)
    - Llama-3.3-70b : modèle open source très performant
    - Vitesse d'inférence extrêmement rapide (LPU Groq)
    - Idéal pour l'extraction structurée de documents
    Pipeline :
        Fichier (PDF/DOCX)
│
▼
        Extraction texte brut (pdfplumber / python-docx)
│
▼
        Nettoyage du texte
│
▼
        ChatGroq (llama-3.3-70b-versatile, temperature=0)
│
▼
        JsonOutputParser + validation Pydantic (CandidateProfile)
│
▼
        dict structuré → JobSearchState
    """
    SYSTEM_PROMPT = """Tu es un expert RH spécialisé dans l'analyse de CVs.
Analyse le texte brut du CV fourni et extrais toutes les informations pertinentes.
Si une information est absente du CV, utilise une chaîne vide "" ou 0 pour les entiers.
Pour le champ "profile_text", génère un paragraphe naturel de 3-5 phrases qui résume
le profil du candidat en incluant ses compétences clés, son expérience et ses objectifs.
Ce texte sera utilisé comme requête pour la recherche d'offres d'emploi.
IMPORTANT : Retourne UNIQUEMENT un objet JSON valide.
Ne commence pas par ```json. N'ajoute aucun texte avant ou après le JSON.
Respecte exactement cette structure :
{{
  "full_name": "",
  "email": "",
  "phone": "",
  "location": "",
  "job_title": "",
  "summary": "",
  "hard_skills": [],
  "soft_skills": [],
  "tools": [],
  "languages": [],
  "education_level": "",
  "experience_years": 0,
  "sectors": [],
  "experiences": [
    {{
      "title": "Intitulé du poste",
      "company": "Nom de l'entreprise",
      "period": "Dates ou durée",
      "description": "Résumé des missions",
      "technologies": ["outil1", "outil2"]
    }}
  ],
  "profile_text": ""
}}
Si aucune expérience professionnelle n'est trouvée, retourne "experiences": [].
IMPORTANT : Extrais TOUTES les expériences professionnelles listées dans le CV. Ne les résume pas trop, garde les détails importants.
"""
    EXTRACTION_PROMPT = """Voici le texte brut extrait du CV :--
{cv_text}--
Extrais toutes les informations et retourne le JSON du profil candidat."""
    def __init__(self, model: str = "llama-3.3-70b-versatile"):
        """
        Args:
            model: Modèle Llama disponible sur Groq.
                   Alternatives : "llama-3.1-8b-instant" (plus rapide, moins précis)
                                  "mixtral-8x7b-32768"   (bon pour les longs CVs)
        """
        self.llm = ChatGroq(
            model=model,
            temperature=0,          # Extraction factuelle → déterministe
            api_key=os.getenv("GROQ_API_KEY"),
            max_tokens=2048,
        )
        self.parser = JsonOutputParser(pydantic_object=CandidateProfile)
        self.prompt = ChatPromptTemplate.from_messages([
            ("system", self.SYSTEM_PROMPT),
            ("human",  self.EXTRACTION_PROMPT),
        ])
        # Chaîne LangChain : Prompt → Llama (Groq) → Parser JSON
        self.chain = self.prompt | self.llm | self.parser
    # ─────────────────────────────────────────
    #  MÉTHODE PRINCIPALE
    # ─────────────────────────────────────────
    def parse(self, file_bytes: bytes, filename: str) -> dict:
        """
        Parse un CV uploadé et retourne un profil structuré.
        Args:
            file_bytes : Contenu binaire du fichier uploadé
            filename   : Nom du fichier (pour détecter le format)
        Returns:
            dict avec tous les champs du profil candidat
        Raises:
            ValueError : Format non supporté ou fichier vide/illisible
        """
        ext = Path(filename).suffix.lower()
        # 1. Extraction du texte brut selon le format
        if ext == ".pdf":
            raw_text = self._extract_from_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            raw_text = self._extract_from_docx(file_bytes)
        else:
            raise ValueError(
                f"Format non supporté : {ext}. Utilisez PDF ou DOCX."
            )
        if not raw_text.strip():
            raise ValueError(
                "Le fichier semble vide ou illisible. "
                "Vérifiez que le CV n'est pas un scan image sans OCR."
            )
        # 2. Nettoyage du texte
        clean_text = self._clean_text(raw_text)
        
        # Log du texte brut pour débogage
        print(f"DEBUG - Texte brut extrait (longueur {len(clean_text)}) : {clean_text[:1000]}...")
        
        # 3. Structuration via Llama (Groq)
        # On limite à 4000 caractères pour plus de contexte
        profile = self.chain.invoke({"cv_text": clean_text[:4000]})
        
        # Log pour débogage
        print(f"DEBUG - Profil extrait : {json.dumps(profile, indent=2)}")
        
        # 4. Ajouter le texte brut pour référence
        profile["raw_cv_text"] = clean_text
        return profile

    def parse_text(self, raw_text: str) -> dict:
        """
        Parse un texte brut (ex: saisi manuellement) et retourne un profil structuré.
        """
        if not raw_text.strip():
            raise ValueError("Le texte fourni est vide.")
            
        clean_text = self._clean_text(raw_text)
        
        print(f"DEBUG - Texte brut extrait (longueur {len(clean_text)}) : {clean_text[:1000]}...")
        
        profile = self.chain.invoke({"cv_text": clean_text[:4000]})
        
        print(f"DEBUG - Profil extrait : {json.dumps(profile, indent=2)}")
        
        profile["raw_cv_text"] = clean_text
        if not profile.get("full_name"):
            profile["full_name"] = "Candidat"
        if not profile.get("job_title"):
            profile["job_title"] = "Profil Saisi"
            
        return profile
    # ─────────────────────────────────────────
    #  EXTRACTEURS DE TEXTE BRUT
    # ─────────────────────────────────────────
    def _extract_from_pdf(self, file_bytes: bytes) -> str:
        """Extrait le texte d'un PDF avec pdfplumber"""
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    def _extract_from_docx(self, file_bytes: bytes) -> str:
        """Extrait le texte d'un fichier DOCX (paragraphes + tableaux)"""
        doc        = DocxDocument(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Les CVs utilisent souvent des tableaux pour la mise en page
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs)
    def _clean_text(self, text: str) -> str:
        """Nettoie le texte extrait du CV"""
        text = re.sub(r'\x00', '', text)         # Null bytes
        text = re.sub(r'\r\n', '\n', text)        # Windows line endings
        text = re.sub(r'\n{3,}', '\n\n', text)    # Lignes vides multiples
        text = re.sub(r'[ \t]{2,}', ' ', text)    # Espaces multiples
        return text.strip()
    #─────────────────────────────────────────
    #  MODE MOCK (sans clé Groq)
    # ─────────────────────────────────────────
    @staticmethod
    def mock_parse(filename: str) -> dict:
        """Retourne un profil fictif pour tester sans clé API"""
        return {
            "full_name":        "Jean-Pierre Mballa",
            "email":            "jpmballa@email.com",
            "phone":            "+237 6XX XXX XXX",
            "location":         "Yaoundé, Cameroun",
            "job_title":        "Ingénieur Data Science",
            "summary":          (
                "Ingénieur Data Science avec 2 ans d'expérience en machine learning "
                "et analyse de données. Passionné par l'IA et les systèmes intelligents. "
                "Cherche un poste challengeant dans le secteur financier ou télécom."
            ),
            "hard_skills":      ["Python", "Machine Learning", "SQL", "Deep Learning", "NLP"],
            "soft_skills":      ["Travail en équipe", "Communication", "Autonomie"],
            "tools":            ["TensorFlow", "Scikit-learn", "Power BI", "Git", "Docker"],
            "languages":        ["Français", "Anglais"],
            "education_level":  "Master en Ingénierie Data Science",
            "experience_years": 2,
            "sectors":          ["Finance", "Télécom", "Data & IA"],
            "experiences": [
                {
                    "title": "Ingénieur IA & Data",
                    "company": "TechCorp Solutions",
                    "period": "2023 - Présent",
                    "description": "Développement et déploiement de modèles de NLP pour l'analyse de sentiments. Optimisation des pipelines de données.",
                    "technologies": ["Python", "Transformers", "PyTorch", "Docker"]
                },
                {
                    "title": "Data Scientist Junior",
                    "company": "CamerData Services",
                    "period": "2021 - 2023",
                    "description": "Analyse prédictive des ventes et création de dashboards interactifs pour la direction.",
                    "technologies": ["Python", "Scikit-learn", "Pandas", "Power BI"]
                }
            ],
            "profile_text":     (
                "Je suis ingénieur Data Science avec 2 ans d'expérience. "
                "Compétences : Python, Machine Learning, SQL, TensorFlow, Power BI. "
                "Je cherche un poste dans l'analyse de données ou le machine learning "
                "à Yaoundé ou Douala dans le secteur financier ou télécom."
            ),
            "raw_cv_text": f"[Texte extrait du fichier : {filename}]"
        }